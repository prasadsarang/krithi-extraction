import numpy as np
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup as BS
import requests
import re
from shutil import copyfile
import os
from docx import Document

# URL = 'http://guru-guha.blogspot.com/'
# r = requests.get(URL)
# soup = BS(r.content,'html.parser')
# f1 = open('Dikshitar_Krithis_links.txt','w')
# for a in soup.find_all('a',href=True)[24::]:
	# link = a['href']
	# print(link)
	# f1.write(link + '\n')
# f1.close()

f1 = open('Dikshitar_Krithis_links.txt','r')
list_of_links = f1.readlines()
for link in list_of_links[185:]:
	print(link)
	r = requests.get(link[:-1])
	soup = BS(r.content,'html.parser')
	title = soup.find('h3',attrs={'itemprop' : 'name'})
	copyfile('C:\\Users\\psarang\\Documents\\Template.docx','C:\\Users\\psarang\\Documents\\Dikshitar_Krithis\\Template.docx')
	os.rename('C:\\Users\\psarang\\Documents\\Dikshitar_Krithis\\Template.docx','C:\\Users\\psarang\\Documents\\Dikshitar_Krithis\\' + title.string[19:-1] + '.docx')
	d = Document('Dikshitar_Krithis\\' +  title.string[19:-1] + '.docx')
	for br in soup.find_all("br"):
		br.replace_with("\n")
	for b in soup.find_all('div',attrs ={'itemprop': 'description articleBody'}):
		if b.text is not None:
			str = b.text.split("Devanagari - Word Division",1)[1]
			str = str.split("Tamil - Word Division",1)[0]
			str1 = str.split("Tamil",1)[0]
			str1 = str1.splitlines()
			str1 = '\n'.join(str1)
			str2 = str.split("Tamil",1)[1]
			str2 = str2.splitlines()
			str2 = '\n'.join(str2[4:-2])
			d.add_paragraph(str2 + '\n\n')
			d.add_paragraph(str1 + '\n')		
	d.save('Dikshitar_Krithis\\' +  title.string[19:-1] + '.docx')
f1.close()